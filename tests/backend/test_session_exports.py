import os
import sys
import uuid
from datetime import datetime, timezone
from io import BytesIO

from fastapi import FastAPI
from fastapi.testclient import TestClient

from docx import Document
from pypdf import PdfReader

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "..", "BE"))

from app.api.endpoints.auth import get_current_user
from app.api.endpoints.sessions import router as sessions_router
from app.db.session import get_db
from app.schemas.user import UserOut
from app.services.session_pdf import build_sessions_pdf
from app.services.session_docx import build_sessions_docx


def _sample_session_payload():
    return {
        "id": "session-123",
        "role_label": "Frontend Developer",
        "level_label": "Junior",
        "status": "COMPLETED",
        "mode": "Camera",
        "created_at_label": "2026-06-18 09:00",
        "completed_at_label": "2026-06-18 09:30",
        "avg_score_label": "8.5/10",
        "question_count": 1,
        "questions": [
            {
                "id": 1,
                "text": "Tell me about the most impactful project in your CV.",
                "category": "Projects",
                "difficulty": "medium",
            }
        ],
        "answers": [
            {
                "question_id": 1,
                "answer_text": "I built a dashboard that improved retention.",
                "feedback": "Structured feedback text.",
                "score": 8.5,
                "telemetry_data": {
                    "gazeRatio": 0.82,
                    "bodyPostureScore": 0.9,
                    "speakingPace": 145,
                    "fillerWordsCount": 3,
                    "presentationConfidence": 78,
                    "blinkRatio": 0.12,
                    "avgTensionScore": 0.34,
                },
                "follow_up_question_text": "Can you explain the trade-offs?",
                "follow_up_answer_text": "I prioritized speed over flexibility at first.",
                "follow_up_feedback": "More detail needed.",
                "follow_up_score": 7.5,
                "follow_up_telemetry_data": {
                    "gazeRatio": 0.75,
                    "bodyPostureScore": 0.88,
                    "speakingPace": 132,
                    "fillerWordsCount": 2,
                    "presentationConfidence": 81,
                    "blinkRatio": 0.1,
                    "avgTensionScore": 0.29,
                },
            }
        ],
    }


def _extract_pdf_text(pdf_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(pdf_bytes))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_docx_text(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    texts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            texts.extend(cell.text for cell in row.cells)
    return "\n".join(texts)


class FakeExportDb:
    async def fetchrow(self, query, *params):
        if "FROM sessions" in query:
            return {
                "id": uuid.UUID("11111111-1111-1111-1111-111111111111"),
                "user_id": uuid.UUID("22222222-2222-2222-2222-222222222222"),
                "major": "technology",
                "role": "frontend",
                "level": "junior",
                "mode": "camera",
                "status": "COMPLETED",
                "created_at": datetime(2026, 6, 18, tzinfo=timezone.utc),
                "completed_at": datetime(2026, 6, 18, tzinfo=timezone.utc),
                "time_limit_minutes": 30,
                "question_count": 1,
                "avg_score": 8.5,
                "evaluation_report": None,
                "practice_plan": None,
            }
        return None

    async def fetch(self, query, *params):
        if "FROM answers" in query:
            return [
                {
                    "id": uuid.UUID("33333333-3333-3333-3333-333333333333"),
                    "session_id": uuid.UUID("11111111-1111-1111-1111-111111111111"),
                    "question_id": 1,
                    "answer_text": "I built a dashboard that improved retention.",
                    "score": 8.5,
                    "feedback": "Structured feedback text.",
                    "telemetry_data": {
                        "gazeRatio": 0.82,
                        "bodyPostureScore": 0.9,
                        "speakingPace": 145,
                        "fillerWordsCount": 3,
                        "presentationConfidence": 78,
                        "blinkRatio": 0.12,
                        "avgTensionScore": 0.34,
                    },
                    "submitted_at": datetime(2026, 6, 18, 9, 10, tzinfo=timezone.utc),
                    "follow_up_id": uuid.UUID("44444444-4444-4444-4444-444444444444"),
                    "follow_up_style": "clarify",
                    "follow_up_question_text": "Can you explain the trade-offs?",
                    "follow_up_answer_text": "I prioritized speed over flexibility at first.",
                    "follow_up_score": 7.5,
                    "follow_up_feedback": "More detail needed.",
                    "follow_up_telemetry_data": {
                        "gazeRatio": 0.75,
                        "bodyPostureScore": 0.88,
                        "speakingPace": 132,
                        "fillerWordsCount": 2,
                        "presentationConfidence": 81,
                        "blinkRatio": 0.1,
                        "avgTensionScore": 0.29,
                    },
                    "follow_up_generated_at": datetime(2026, 6, 18, 9, 11, tzinfo=timezone.utc),
                    "follow_up_answered_at": datetime(2026, 6, 18, 9, 12, tzinfo=timezone.utc),
                }
            ]
        if "FROM session_question_sets" in query:
            return [
                {
                    "id": 1,
                    "major": "technology",
                    "role": "frontend",
                    "level": "junior",
                    "text": "Tell me about the most impactful project in your CV.",
                    "text_en": "Tell me about the most impactful project in your CV.",
                    "text_vi": "Hãy kể về dự án có tác động lớn nhất trong CV của bạn.",
                    "category": "Projects",
                    "category_en": "Projects",
                    "category_vi": "Dự án",
                    "difficulty": "medium",
                    "tags": ["CV-based"],
                }
            ]
        return []


def test_build_sessions_pdf_includes_video_metrics_section():
    pdf_bytes = build_sessions_pdf(sessions=[_sample_session_payload()], language="en", export_all=False)
    text = _extract_pdf_text(pdf_bytes)

    assert "Video metrics" in text
    assert "Eye contact" in text
    assert "Speaking pace" in text
    assert "Follow-up video metrics" in text


def test_build_sessions_docx_includes_video_metrics_section():
    docx_bytes = build_sessions_docx(sessions=[_sample_session_payload()], language="en", export_all=False)
    text = _extract_docx_text(docx_bytes)

    assert "Video metrics" in text
    assert "Eye contact" in text
    assert "Speaking pace" in text
    assert "Follow-up video metrics" in text


def test_export_session_docx_endpoint_returns_docx_with_video_metrics(monkeypatch):
    import app.api.endpoints.sessions as sessions_module

    app = FastAPI()
    app.include_router(sessions_router, prefix="/api/sessions")

    fake_db = FakeExportDb()

    async def override_db():
        yield fake_db

    async def override_current_user():
        return UserOut(
            id=uuid.UUID("22222222-2222-2222-2222-222222222222"),
            email="candidate@example.com",
            created_at=datetime(2026, 6, 18, tzinfo=timezone.utc),
            full_name="Candidate",
            is_admin=False,
        )

    async def fake_plan_snapshot(db, user_id):
        return {
            "is_admin": False,
            "plan_tier": "pro",
            "plan_status": "active",
        }

    monkeypatch.setattr(sessions_module, "get_user_plan_snapshot", fake_plan_snapshot)

    app.dependency_overrides[get_db] = override_db
    app.dependency_overrides[get_current_user] = override_current_user

    client = TestClient(app)
    response = client.get("/api/sessions/11111111-1111-1111-1111-111111111111/export-docx")

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert response.headers["content-disposition"].endswith('.docx"')
    text = _extract_docx_text(response.content)
    assert "Video metrics" in text
    assert "Follow-up video metrics" in text
