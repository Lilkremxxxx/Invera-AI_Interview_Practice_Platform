import docx

doc = docx.Document("/home/nhatbang/EXE101/PRJ/docs/Crawl_qst.docx")

# Let's print paragraphs 45 to 80
for idx in range(45, 80):
    print(f"{idx} [{doc.paragraphs[idx].style.name}]: {doc.paragraphs[idx].text}")
