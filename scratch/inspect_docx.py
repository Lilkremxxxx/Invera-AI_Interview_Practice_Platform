import docx

doc = docx.Document("/home/nhatbang/EXE101/PRJ/docs/Crawl_qst.docx")
print("Number of paragraphs:", len(doc.paragraphs))

# Let's inspect paragraphs starting from index 100 to 150
print("\n--- Paragraphs 100 to 160 ---")
for idx in range(100, 160):
    text = doc.paragraphs[idx].text.strip()
    if text:
        print(f"{idx}: {text[:150]}")

# Let's also inspect paragraphs starting from index 500 to 550
print("\n--- Paragraphs 500 to 550 ---")
for idx in range(500, 550):
    text = doc.paragraphs[idx].text.strip()
    if text:
        print(f"{idx}: {text[:150]}")
