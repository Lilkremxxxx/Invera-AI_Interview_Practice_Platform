import os
import re
from docx import Document

DOCX_PATH = "/home/nhatbang/EXE101/PRJ/docs/Crawl_qst.docx"

def inspect():
    if not os.path.exists(DOCX_PATH):
        print(f"File not found: {DOCX_PATH}")
        return
        
    doc = Document(DOCX_PATH)
    print(f"Document paragraphs count: {len(doc.paragraphs)}")
    
    # We want to count:
    # 1. Heading 2 (Positions)
    # 2. Heading 3 (Levels)
    # 3. Questions (starts with 'Câu ')
    # 4. Answers (starts with 'Trả lời: ')
    # 5. Tags (starts with 'Tags: ')
    
    positions = []
    levels = []
    questions_count = 0
    answers_count = 0
    tags_count = 0
    
    current_pos = None
    current_level = None
    
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style_name = para.style.name
        
        if style_name.startswith('Heading 2'):
            positions.append(text)
            current_pos = text
        elif style_name.startswith('Heading 3'):
            levels.append(text)
            current_level = text
        elif text.startswith('Câu '):
            questions_count += 1
        elif text.startswith('Trả lời: '):
            answers_count += 1
        elif text.startswith('Tags: '):
            tags_count += 1
            
    print("\nExtraction Summary:")
    print(f"  Positions found ({len(positions)}): {positions}")
    print(f"  Levels subheadings found: {len(levels)}")
    print(f"  Questions found: {questions_count}")
    print(f"  Answers found: {answers_count}")
    print(f"  Tags found: {tags_count}")

if __name__ == "__main__":
    inspect()
