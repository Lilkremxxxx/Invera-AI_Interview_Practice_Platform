#!/usr/bin/env python3
"""Compile JSON interview data (questions + answers + tags) into a DOCX file."""

import json
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DATA_DIR = "/home/nhatbang/EXE101/PRJ/docs/crawl_data"
OUTPUT_PATH = "/home/nhatbang/EXE101/PRJ/docs/Crawl_qst.docx"

POSITION_FILES = [
    ("frontend_developer.json", "Frontend Developer", "Tech"),
    ("backend_developer.json", "Backend Developer", "Tech"),
    ("full_stack_developer.json", "Full Stack Developer", "Tech"),
    ("data_scientist.json", "Data Scientist", "Tech"),
    ("machine_learning_engineer.json", "Machine Learning Engineer", "Tech"),
    ("devops_engineer.json", "DevOps Engineer", "Tech"),
    ("product_manager.json", "Product Manager", "Tech"),
    ("ux_designer.json", "UX Designer", "Tech"),
    ("business_analyst.json", "Business Analyst", "Business"),
    ("operations_analyst.json", "Operations Analyst", "Business"),
    ("sales_representative.json", "Sales Representative", "Business"),
    ("marketing_specialist.json", "Marketing Specialist", "Business"),
    ("marketing_manager.json", "Marketing Manager", "Business"),
    ("financial_analyst.json", "Financial Analyst", "Finance"),
    ("accountant.json", "Accountant", "Finance"),
    ("auditor.json", "Auditor", "Finance"),
    ("investment_banking_analyst.json", "Investment Banking Analyst", "Finance"),
]

LEVELS = ["Intern", "Fresher", "Junior", "Middle", "Senior"]

doc = Document()

# --- Global Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# ====== HELPER: add formatted paragraph ======
def add_para(text, bold=False, size=11, color=None, align=None, space_after=4, space_before=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

def add_rich_para(parts, space_after=4, space_before=2):
    """parts = list of (text, bold, size, color)"""
    p = doc.add_paragraph()
    for text, bold, size, color in parts:
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p

# Load all data
all_data = {}
grand_total = 0
category_totals = {"Tech": 0, "Business": 0, "Finance": 0}

for filename, pos_name, category in POSITION_FILES:
    path = os.path.join(DATA_DIR, filename)
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        all_data[pos_name] = data
        pos_total = sum(len(data.get("questions", {}).get(level, [])) for level in LEVELS)
        grand_total += pos_total
        category_totals[category] += pos_total
    except Exception as e:
        print(f"Error loading {filename}: {e}")

# ====== COVER PAGE ======
for _ in range(6):
    doc.add_paragraph()

add_para("TUY\u1ec2N T\u1eacP C\u00c2U H\u1eceI PH\u1eceNG V\u1ea4N", 
         bold=True, size=28, color=RGBColor(0x1A, 0x3C, 0x6E), 
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Theo C\u1ea5p \u0110\u1ed9: Intern - Fresher - Junior - Middle - Senior",
         size=16, color=RGBColor(0x33, 0x66, 0x99),
         align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
add_para("17 V\u1ecb tr\u00ed thu\u1ed9c 3 Nh\u00f3m ng\u00e0nh: Tech | Business | Finance",
         size=13, color=RGBColor(0x66, 0x66, 0x66),
         align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
doc.add_paragraph()

# ====== TABLE OF CONTENTS ======
doc.add_page_break()
doc.add_heading("M\u1ee4C L\u1ee4C", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

for cat_title, cat_key in [
    ("NH\u00d3M TECH", "Tech"),
    ("NH\u00d3M BUSINESS", "Business"),
    ("NH\u00d3M FINANCE", "Finance"),
]:
    doc.add_heading(cat_title, level=2)
    for pos_file, pos_name, category in POSITION_FILES:
        if category == cat_key:
            data = all_data.get(pos_name, {})
            total_q = sum(len(data.get("questions", {}).get(level, [])) for level in LEVELS)
            add_para(f"      {pos_name} ...................................................... {total_q} cau")
    doc.add_paragraph()

# ====== CONTENT ======
BLUE = RGBColor(0x1A, 0x3C, 0x6E)
LIGHT_BLUE = RGBColor(0x33, 0x66, 0x99)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK_GREEN = RGBColor(0x1B, 0x5E, 0x20)
TAG_COLOR = RGBColor(0x55, 0x7B, 0xC7)

for cat_title, cat_key in [("NHÓM TECH", "Tech"), ("NHÓM BUSINESS", "Business"), ("NHÓM FINANCE", "Finance")]:
    doc.add_page_break()
    doc.add_heading(cat_title, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(f"Tổng số câu hỏi: {category_totals[cat_key]}", size=12, color=LIGHT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    for level in LEVELS:
        level_total = sum(
            len(all_data.get(pos_name, {}).get("questions", {}).get(level, []))
            for pos_file, pos_name, category in POSITION_FILES if category == cat_key
        )
        add_para(f"Cấp độ {level}: {level_total} câu hỏi", size=11, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    doc.add_page_break()
    
    for pos_file, pos_name, category in POSITION_FILES:
        if category != cat_key:
            continue
        
        data = all_data.get(pos_name, {})
        if not data:
            continue
        
        doc.add_heading(pos_name, level=2)
        
        for level in LEVELS:
            questions = data.get("questions", {}).get(level, [])
            answers = data.get("answers", {}).get(level, [])
            tags_list = data.get("tags", {}).get(level, [])
            
            if not questions:
                continue
            
            doc.add_heading(f"Cấp độ: {level} ({len(questions)} câu hỏi)", level=3)
            
            for idx, q in enumerate(questions, 1):
                # Question
                add_rich_para([
                    (f"Câu {idx}: ", True, 11, BLUE),
                    (q, False, 11, None),
                ], space_before=4, space_after=2)
                
                # Answer
                if idx - 1 < len(answers):
                    ans = answers[idx - 1]
                    # Check if answer is detailed or generic
                    add_rich_para([
                        ("Trả lời: ", True, 10, DARK_GREEN),
                        (ans, False, 10, None),
                    ], space_after=2)
                
                # Tags
                if idx - 1 < len(tags_list):
                    tags = tags_list[idx - 1]
                    if isinstance(tags, list):
                        tags_text = " | ".join(tags)
                        add_rich_para([
                            ("Tags: ", False, 9, TAG_COLOR),
                            (tags_text, False, 9, TAG_COLOR),
                        ], space_after=4)
                
                # Separator between Q&A pairs
                add_para("─" * 60, size=8, color=GRAY, space_after=4)
            
            doc.add_paragraph()  # spacing between levels
        
        doc.add_page_break()

# ====== STATISTICS PAGE ======
doc.add_heading("THỐNG KÊ TỔNG QUAN", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

table = doc.add_table(rows=1, cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
for i, h in enumerate(["Vị trí", "Intern", "Fresher", "Junior", "Middle", "Senior", "Tổng"]):
    hdr[i].text = h
    for paragraph in hdr[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

grand = {"Intern": 0, "Fresher": 0, "Junior": 0, "Middle": 0, "Senior": 0}
grand_row_total = 0

for pos_file, pos_name, category in POSITION_FILES:
    data = all_data.get(pos_name, {})
    row = table.add_row()
    cells = row.cells
    cells[0].text = pos_name
    row_total = 0
    for i, level in enumerate(LEVELS):
        count = len(data.get("questions", {}).get(level, []))
        cells[i+1].text = str(count)
        grand[level] += count
        row_total += count
    cells[6].text = str(row_total)
    grand_row_total += row_total

# Grand total row
row = table.add_row()
cells = row.cells
cells[0].text = "TỔNG CỘNG"
for r in cells[0].paragraphs[0].runs:
    r.bold = True
for i, level in enumerate(LEVELS):
    cells[i+1].text = str(grand[level])
    for r in cells[i+1].paragraphs[0].runs:
        r.bold = True
    grand_row_total += 0  # already counted above
cells[6].text = str(grand_row_total)
for r in cells[6].paragraphs[0].runs:
    r.bold = True

doc.add_paragraph()
add_para(f"Tổng số câu hỏi: {grand_total}", bold=True, size=16, color=BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_paragraph()
cat_text = "  ".join(f"{cat}: {total} câu" for cat, total in category_totals.items())
add_para(cat_text, size=12, color=LIGHT_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"SUCCESS: Document saved to {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH)} bytes")
print(f"Total questions with answers and tags: {grand_total}")
