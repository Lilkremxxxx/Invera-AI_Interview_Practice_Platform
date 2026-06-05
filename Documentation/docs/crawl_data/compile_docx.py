import json
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

DATA_DIR = "/home/nhatbang/EXE101/PRJ/docs/crawl_data"
OUTPUT_PATH = "/home/nhatbang/EXE101/PRJ/docs/Crawl_qst.docx"

POSITION_FILES = [
    ("frontend_developer.json", "Frontend Developer", "Tech"),
    ("backend_developer.json", "Backend Developer", "Tech"),
    ("full_stack_developer.json", "Full Stack Developer", "Tech"),
    ("data_scientist.json", "Data Scientist", "Tech"),
    ("machine_learning_engineer.json", "Machine Learning Engineer", "Tech"),
    ("devops_engineer.json", "DevOps Engineer", "Tech"),
    ("product_manager.json", "Product Manager", "Tech"),
    ("ux_designer.json", "UX Designer", "Tech"),
    ("business_analyst.json", "Business Analyst", "Business"),
    ("operations_analyst.json", "Operations Analyst", "Business"),
    ("sales_representative.json", "Sales Representative", "Business"),
    ("marketing_specialist.json", "Marketing Specialist", "Business"),
    ("marketing_manager.json", "Marketing Manager", "Business"),
    ("financial_analyst.json", "Financial Analyst", "Finance"),
    ("accountant.json", "Accountant", "Finance"),
    ("auditor.json", "Auditor", "Finance"),
    ("investment_banking_analyst.json", "Investment Banking Analyst", "Finance"),
]

LEVELS = ["Intern", "Fresher", "Junior", "Middle", "Senior"]

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# ====== COVER PAGE ======
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("TUYỂN TẬP CÂU HỎI PHỎNG VẤN")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Theo Cấp Độ: Intern - Fresher - Junior - Middle - Senior")
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run("17 Vị trí thuộc 3 Nhóm ngành: Tech | Business | Finance")
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
doc.add_paragraph()

# Table of counts
counts_para = doc.add_paragraph()
counts_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Load all data to compute totals
all_data = {}
grand_total = 0
category_totals = {"Tech": 0, "Business": 0, "Finance": 0}

for filename, pos_name, category in POSITION_FILES:
    path = os.path.join(DATA_DIR, filename)
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        all_data[pos_name] = data
        pos_total = sum(len(data.get("questions", {}).get(level, [])) for level in LEVELS)
        grand_total += pos_total
        category_totals[category] += pos_total
    except Exception as e:
        print(f"Error loading {filename}: {e}")

for _ in range(3):
    doc.add_paragraph()

# ====== TABLE OF CONTENTS ======
doc.add_page_break()
toc_title = doc.add_heading("MỤC LỤC", level=1)
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

categories_order = [
    ("NHÓM TECH", "Tech"),
    ("NHÓM BUSINESS", "Business"),
    ("NHÓM FINANCE", "Finance"),
]

for cat_title, cat_key in categories_order:
    heading = doc.add_heading(cat_title, level=2)
    heading_num = doc.add_paragraph()
    for pos_file, pos_name, category in POSITION_FILES:
        if category == cat_key:
            data = all_data.get(pos_name, {})
            total_q = sum(len(data.get("questions", {}).get(level, [])) for level in LEVELS)
            p = doc.add_paragraph(
                f"      {pos_name} ...................................................... {total_q} câu",
                style='List Bullet'
            )
    doc.add_paragraph()

# ====== CONTENT ======
doc.add_page_break()
for cat_title, cat_key in categories_order:
    # Category divider page
    cat_heading = doc.add_heading(cat_title, level=1)
    cat_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cat_info = doc.add_paragraph()
    cat_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cat_info.add_run(f"Tổng số câu hỏi: {category_totals[cat_key]}")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
    
    for level in LEVELS:
        level_total = 0
        for pos_file, pos_name, category in POSITION_FILES:
            if category == cat_key:
                data = all_data.get(pos_name, {})
                level_total += len(data.get("questions", {}).get(level, []))
        
        level_para = doc.add_paragraph()
        level_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = level_para.add_run(f"Cấp độ {level}: {level_total} câu hỏi")
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.add_page_break()
    
    for pos_file, pos_name, category in POSITION_FILES:
        if category != cat_key:
            continue
        
        data = all_data.get(pos_name, {})
        if not data:
            continue
        
        # Position title page
        pos_heading = doc.add_heading(pos_name, level=2)
        
        for level in LEVELS:
            questions = data.get("questions", {}).get(level, [])
            if not questions:
                continue
            
            level_title = doc.add_heading(f"Cấp độ: {level} ({len(questions)} câu hỏi)", level=3)
            
            for idx, q in enumerate(questions, 1):
                p = doc.add_paragraph()
                run_num = p.add_run(f"Câu {idx}: ")
                run_num.bold = True
                run_num.font.size = Pt(11)
                run_q = p.add_run(q)
                run_q.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.space_before = Pt(2)
            
            doc.add_paragraph()  # spacing between levels
        
        doc.add_page_break()

# ====== STATISTICS PAGE ======
stats_title = doc.add_heading("THỐNG KÊ TỔNG QUAN", level=1)
stats_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# Summary table
table = doc.add_table(rows=1, cols=7)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
headers = ["Vị trí", "Intern", "Fresher", "Junior", "Middle", "Senior", "Tổng"]
for i, h in enumerate(headers):
    hdr[i].text = h
    for paragraph in hdr[i].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(9)

grand = {"Intern": 0, "Fresher": 0, "Junior": 0, "Middle": 0, "Senior": 0}

for pos_file, pos_name, category in POSITION_FILES:
    data = all_data.get(pos_name, {})
    row = table.add_row()
    cells = row.cells
    cells[0].text = pos_name
    row_total = 0
    for i, level in enumerate(LEVELS):
        count = len(data.get("questions", {}).get(level, []))
        cells[i+1].text = str(count)
        grand[level] += count
        row_total += count
    cells[6].text = str(row_total)
    for cell in cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)

# Grand total row
row = table.add_row()
cells = row.cells
cells[0].text = "TỔNG CỘNG"
cells[0].paragraphs[0].runs[0].bold = True
grand_row_total = 0
for i, level in enumerate(LEVELS):
    cells[i+1].text = str(grand[level])
    cells[i+1].paragraphs[0].runs[0].bold = True
    grand_row_total += grand[level]
cells[6].text = str(grand_row_total)
cells[6].paragraphs[0].runs[0].bold = True

for cell in cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

doc.add_paragraph()
summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = summary.add_run(f"Tổng số câu hỏi: {grand_total}")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

doc.add_paragraph()
cat_summary = doc.add_paragraph()
cat_summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
for cat, total in category_totals.items():
    run = cat_summary.add_run(f"  {cat}: {total} câu  |")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)

# Save
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"SUCCESS: Document saved to {OUTPUT_PATH}")
print(f"Total questions: {grand_total}")
print(f"Size: {os.path.getsize(OUTPUT_PATH)} bytes")
