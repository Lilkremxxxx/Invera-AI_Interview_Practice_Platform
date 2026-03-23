from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi import UploadFile


DOCX_MAX_BYTES = 10 * 1024 * 1024
ALLOWED_DOCX_EXTENSIONS = {".docx"}
ALLOWED_DOCX_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
}


class QnaDocxValidationError(ValueError):
    pass


def _normalized_suffix(filename: str | None) -> str:
    return Path(filename or "").suffix.lower()


def _read_with_limit(upload: UploadFile, max_bytes: int) -> bytes:
    upload.file.seek(0)
    data = upload.file.read(max_bytes + 1)
    upload.file.seek(0)
    if len(data) > max_bytes:
        raise QnaDocxValidationError(f"DOCX file exceeds the {max_bytes // (1024 * 1024)}MB limit.")
    return data


def extract_docx_text(upload: UploadFile) -> tuple[str, str]:
    suffix = _normalized_suffix(upload.filename)
    if suffix not in ALLOWED_DOCX_EXTENSIONS:
        raise QnaDocxValidationError("Only DOCX files are supported right now.")

    content_type = (upload.content_type or "").lower()
    if content_type and content_type not in ALLOWED_DOCX_CONTENT_TYPES:
        raise QnaDocxValidationError("Unsupported DOCX content type.")

    raw = _read_with_limit(upload, DOCX_MAX_BYTES)
    try:
        document = Document(BytesIO(raw))
    except Exception as exc:  # pragma: no cover - parser-specific
        raise QnaDocxValidationError("Unable to read the DOCX file.") from exc

    chunks: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chunks.append(row_text)

    extracted = "\n".join(chunks).strip()
    if not extracted:
        raise QnaDocxValidationError("The DOCX file does not contain readable text.")

    return upload.filename or "document.docx", extracted[:20000]
