from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.core.text_processing import sanitize_user_text
from app.services.session_export import build_video_metric_blocks, sanitize_metric_note


def _text_labels(language: str) -> dict[str, str]:
    if language == "vi":
        return {
            "report_title": "Invera Session Export",
            "report_subtitle": "Bản xuất câu hỏi, câu trả lời và góp ý phỏng vấn",
            "session": "Session",
            "role": "Vai trò",
            "level": "Cấp độ",
            "status": "Trạng thái",
            "mode": "Chế độ",
            "created": "Tạo lúc",
            "completed": "Hoàn thành lúc",
            "avg_score": "Điểm trung bình",
            "question_count": "Số câu hỏi",
            "question": "Question",
            "category": "Chủ đề",
            "difficulty": "Độ khó",
            "user_answer": "Câu trả lời của user",
            "feedback": "Gợi ý và feedback",
            "score": "Điểm",
            "empty_answer": "Chưa có câu trả lời cho mục này.",
            "empty_feedback": "Chưa có feedback.",
            "in_progress": "Đang làm",
            "completed_status": "Hoàn thành",
            "unknown": "Chưa có",
            "ai_evaluation": "Báo cáo đánh giá chi tiết từ AI (Giao tiếp & Telemetry)",
            "practice_plan": "Kế hoạch luyện tập đề xuất từ AI",
        }
    return {
        "report_title": "Invera Session Export",
        "report_subtitle": "Interview questions, answers, and coaching feedback",
        "session": "Session",
        "role": "Role",
        "level": "Level",
        "status": "Status",
        "mode": "Mode",
        "created": "Created",
        "completed": "Completed",
        "avg_score": "Average score",
        "question_count": "Question count",
        "question": "Question",
        "category": "Category",
        "difficulty": "Difficulty",
        "user_answer": "User answer",
        "feedback": "Feedback and suggestions",
        "score": "Score",
        "empty_answer": "No answer was submitted for this item yet.",
        "empty_feedback": "No feedback yet.",
        "in_progress": "In progress",
        "completed_status": "Completed",
        "unknown": "Not available",
        "ai_evaluation": "Detailed AI Evaluation Report (Communication & Telemetry)",
        "practice_plan": "AI Recommended Practice Plan",
    }


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10, color: str = "1f2937") -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Noto Sans"


def _style_document(doc: Document) -> None:
    styles = doc.styles
    for style_name in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        if style_name in styles:
            styles[style_name].font.name = "Noto Sans"
    if "Normal" in styles:
        styles["Normal"].font.size = Pt(10)


def _add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = "Noto Sans"
    run.font.color.rgb = RGBColor.from_string("0f172a")


def _add_subtitle(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Noto Sans"
    run.font.color.rgb = RGBColor.from_string("475569")


def _session_meta_table(doc: Document, session: dict[str, Any], labels: dict[str, str]) -> None:
    def value(raw: Any) -> str:
        if raw is None or raw == "":
            return labels["unknown"]
        return str(raw)

    rows = [
        (labels["role"], value(session.get("role_label"))),
        (labels["level"], value(session.get("level_label"))),
        (labels["status"], labels["completed_status"] if value(session.get("status")) == "COMPLETED" else labels["in_progress"]),
        (labels["mode"], value(session.get("mode"))),
        (labels["created"], value(session.get("created_at_label"))),
        (labels["completed"], value(session.get("completed_at_label"))),
        (labels["avg_score"], value(session.get("avg_score_label"))),
        (labels["question_count"], value(session.get("question_count"))),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for row_idx, (label, value_text) in enumerate(rows):
        _set_cell_text(table.rows[row_idx].cells[0], label, bold=True, size=9, color="0f172a")
        _set_cell_text(table.rows[row_idx].cells[1], value_text, size=9, color="334155")
        _set_cell_shading(table.rows[row_idx].cells[0], "eff6ff")
        _set_cell_shading(table.rows[row_idx].cells[1], "f8fafc")


def _metric_table(doc: Document, title: str, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=max(1, len(rows) + 1), cols=2)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].merge(header[1])
    _set_cell_text(header[0], title, bold=True, size=10, color="0f172a")
    _set_cell_shading(header[0], "dbeafe")

    if not rows:
      _set_cell_text(table.rows[1].cells[0], "No data", size=9)

    for row_idx, (label, value) in enumerate(rows, start=1):
        _set_cell_text(table.rows[row_idx].cells[0], label, bold=True, size=9, color="0f172a")
        _set_cell_text(table.rows[row_idx].cells[1], value, size=9, color="334155")
        _set_cell_shading(table.rows[row_idx].cells[0], "f8fafc")
        _set_cell_shading(table.rows[row_idx].cells[1], "ffffff")


def _add_paragraph(doc: Document, text: str, *, bold: bool = False, size: int = 10, color: str = "1f2937") -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Noto Sans"


def _add_formatted_text_to_paragraph(paragraph, text: str) -> None:
    import re
    # Split the string by bold and italic Markdown sequences
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|_.*?_)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            run.font.name = "Noto Sans"
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.font.italic = True
            run.font.name = "Noto Sans"
        else:
            run = paragraph.add_run(part)
            run.font.name = "Noto Sans"


def add_markdown_paragraphs_docx(doc: Document, md_text: str) -> None:
    if not md_text:
        return

    lines = md_text.split("\n")
    for line in lines:
        line_str = line.strip()
        if not line_str:
            continue

        if line_str.startswith("###") or line_str.startswith("##") or line_str.startswith("#"):
            text = line_str.lstrip("#").strip()
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Noto Sans"
            run.font.color.rgb = RGBColor.from_string("0f172a")
            paragraph.paragraph_format.space_before = Pt(8)
            paragraph.paragraph_format.space_after = Pt(4)
        elif line_str.startswith("- ") or line_str.startswith("* "):
            text = line_str[2:].strip()
            paragraph = doc.add_paragraph(style='List Bullet')
            _add_formatted_text_to_paragraph(paragraph, text)
            paragraph.paragraph_format.space_after = Pt(3)
        elif line_str.startswith(">"):
            text = line_str[1:].strip()
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.25)
            run = paragraph.add_run(text)
            run.font.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor.from_string("475569")
            paragraph.paragraph_format.space_after = Pt(4)
        else:
            paragraph = doc.add_paragraph()
            _add_formatted_text_to_paragraph(paragraph, line_str)
            paragraph.paragraph_format.space_after = Pt(4)


def build_sessions_docx(*, sessions: list[dict[str, Any]], language: str, export_all: bool) -> bytes:
    labels = _text_labels(language)
    doc = Document()
    _style_document(doc)

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    _add_heading(doc, labels["report_title"])
    _add_subtitle(doc, labels["report_subtitle"])

    for session_index, session in enumerate(sessions, start=1):
        if session_index > 1:
            doc.add_page_break()

        title = f"{labels['session']} {session_index}"
        if session.get("role_label"):
            title = f"{title}: {sanitize_user_text(session['role_label'])}"
        _add_paragraph(doc, title, bold=True, size=16, color="0f172a")
        _session_meta_table(doc, session, labels)
        doc.add_paragraph("")

        # Append AI evaluation report if present
        eval_report = session.get("evaluation_report")
        if eval_report:
            _add_paragraph(doc, labels["ai_evaluation"], bold=True, size=11, color="0f766e")
            add_markdown_paragraphs_docx(doc, eval_report)
            doc.add_paragraph("")

        # Append proposed practice plan if present
        practice_plan = session.get("practice_plan")
        if practice_plan:
            _add_paragraph(doc, labels["practice_plan"], bold=True, size=11, color="0f766e")
            add_markdown_paragraphs_docx(doc, practice_plan)
            doc.add_paragraph("")

        questions = session.get("questions") or []
        answer_map = {answer["question_id"]: answer for answer in session.get("answers", [])}

        for question_index, question in enumerate(questions, start=1):
            answer = answer_map.get(question["id"])
            _add_paragraph(doc, f"{labels['question']} {question_index}", bold=True, size=12, color="111827")
            meta_bits = [question.get("category"), question.get("difficulty")]
            meta_line = " • ".join(bit for bit in meta_bits if bit)
            if meta_line:
                _add_paragraph(doc, meta_line, size=9, color="475569")
            _add_paragraph(doc, sanitize_user_text(question.get("text", "")), size=10)

            _add_paragraph(doc, labels["user_answer"], bold=True, size=10, color="0f766e")
            _add_paragraph(
                doc,
                sanitize_user_text(answer["answer_text"]) if answer and answer.get("answer_text") else labels["empty_answer"],
                size=10,
            )

            _add_paragraph(doc, labels["feedback"], bold=True, size=10, color="0f766e")
            if answer and answer.get("score") is not None:
                _add_paragraph(doc, f"{labels['score']}: {answer['score']}/10", size=10)
            feedback_text = sanitize_user_text(answer["feedback"]) if answer and answer.get("feedback") else labels["empty_feedback"]
            _add_paragraph(doc, feedback_text, size=10)

            metric_blocks = build_video_metric_blocks(answer or {}, language) if answer else []
            if metric_blocks:
                for metric_block in metric_blocks:
                    _metric_table(doc, metric_block["title"], metric_block["rows"])
            else:
                _add_paragraph(doc, sanitize_metric_note(None, language), size=9, color="64748b")

            doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_session_docx_filename(role: str, session_id: str, export_all: bool) -> str:
    if export_all:
        return "invera-sessions-export.docx"
    slug = sanitize_user_text(role).lower().replace(" ", "-") or "session"
    return f"invera-session-{slug}-{session_id[:8]}.docx"
